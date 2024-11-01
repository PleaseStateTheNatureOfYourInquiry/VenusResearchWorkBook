
# Standard imports.
import os
import numpy as np

# Special imports.
from docx import Document
from docx.shared import Pt


# Custom imports.
from HandyTools import HandyTools


workBookPath = '/Users/maarten/Science/Venus/VenusResearchWorkBook/Temperature-UVBrightness-Project'

tableContentImagesPerOrbit= HandyTools.readTable ( os.path.join (workBookPath, 'VMC/Step01/VMCSelectedImages_orbits_later_than_1188.dat') )
tableContentRFRPerOrbit = HandyTools.readTable ( os.path.join (workBookPath, 'VMC/Step03/RadianceFactorRatiosPerOrbit.dat') )

numberOfRows = len ( np.where ( tableContentRFRPerOrbit [0][0] > 1188 ) [0] )
numberOfColumns = 8

tableDocument = Document()
run = tableDocument.add_paragraph ().add_run ()
font = run.font
font.name = 'Times New Roman'
font.size = Pt (12)


table = tableDocument.add_table (rows = numberOfRows + 1, cols = numberOfColumns)

iRowAdded = 0
# table.cell (iRowAdded,0).text = 'VEX OrbitID'
# table.cell (iRowAdded,1).text = 'Date'
# table.cell (iRowAdded,2).text = 'VeRa\nϴ, φ'
# table.cell (iRowAdded,3).text = '# VMC images'
# table.cell (iRowAdded,4).text = 'RFR\n(average)'
# table.cell (iRowAdded,5).text = '∆RFR'
# table.cell (iRowAdded,6).text = 'RFR\n(median)'
# table.cell (iRowAdded,6).text = '∆RFR'

table.cell (iRowAdded,0).add_paragraph (text='VEX OrbitID', style = 'Body Text')
table.cell (iRowAdded,1).add_paragraph (text='Date', style = 'Body Text')
table.cell (iRowAdded,2).add_paragraph (text='VeRa\nϴ, φ', style = 'Body Text')
table.cell (iRowAdded,3).add_paragraph (text='# VMC images', style = 'Body Text')
table.cell (iRowAdded,4).add_paragraph (text='RFR\n(average)', style = 'Body Text')
table.cell (iRowAdded,5).add_paragraph (text='∆RFR', style = 'Body Text')
table.cell (iRowAdded,6).add_paragraph (text='RFR\n(median)', style = 'Body Text')
table.cell (iRowAdded,7).add_paragraph (text='∆RFR', style = 'Body Text')

for iRow in range ( len (tableContentRFRPerOrbit [0][0]) ):

    if tableContentRFRPerOrbit [1][0][iRow] > '1188':
        
        iRowAdded += 1
        
        iRowImageTable = 0
        while tableContentImagesPerOrbit [1][0][iRowImageTable] != tableContentRFRPerOrbit [1][0][iRow]:
        
            iRowImageTable += 1
                
        if float ( tableContentImagesPerOrbit [1][6][iRowImageTable] ) < 0:
        
            latitude = int ( float ( tableContentImagesPerOrbit [1][6][iRowImageTable] ) - 0.5 )
        
        longitude = int ( float ( tableContentImagesPerOrbit [1][7][iRowImageTable] ) + 0.5)    
            
        row_cells = table.rows [iRowAdded].cells
        row_cells [0].text = tableContentRFRPerOrbit [1][0][iRow]
        row_cells [1].text = tableContentImagesPerOrbit [1][2][iRowImageTable]
        row_cells [2].text = '{:03d}˚, {:03d}˚'.format (latitude, longitude)
        row_cells [3].text = tableContentRFRPerOrbit [1][6][iRow]
        row_cells [4].text = tableContentRFRPerOrbit [1][2][iRow]
        row_cells [5].text = tableContentRFRPerOrbit [1][3][iRow]
        row_cells [6].text = tableContentRFRPerOrbit [1][4][iRow]
        row_cells [7].text = tableContentRFRPerOrbit [1][5][iRow]



tableDocument.save('2024_Roos-Serote_Table01_Unformatted.docx')

